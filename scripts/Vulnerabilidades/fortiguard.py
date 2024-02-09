## BUGS
# - Not working yet my friend

import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime

# Set the URL
url = "https://www.fortiguard.com/psirt"

try:
    # Send a GET request to the URL
    response = requests.get(url)
    
    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(response.content, "html.parser")
        
        # Find all rows containing vulnerability information
        vulnerability_rows = soup.find_all("div", class_="row")
        
        if not vulnerability_rows:
            print("No vulnerabilities found on the page.")
        
        # Create a Document object to store the content
        doc = Document()
        
        # Get the current date
        current_date = datetime.now().strftime('%d-%m-%Y')
        
        # Initialize vulnerability count
        vulnerability_count = 0
        
        # Loop through the vulnerability rows and extract the information
        for row in vulnerability_rows:
            # Find the link to the vulnerability detail page
            link_element = row.find("a", class_="link")
            if link_element:
                vulnerability_url = "https://www.fortiguard.com" + link_element["href"]
                
                # Send a GET request to the vulnerability detail page
                vulnerability_response = requests.get(vulnerability_url)
                
                # Check if the request was successful (status code 200)
                if vulnerability_response.status_code == 200:
                    # Parse the HTML content of the vulnerability detail page
                    vulnerability_soup = BeautifulSoup(vulnerability_response.content, "html.parser")
                    
                    # Extract data from the vulnerability detail page
                    ir_number = vulnerability_soup.find("td", text="IR Number").find_next_sibling("td").text.strip()
                    date = vulnerability_soup.find("td", text="Date").find_next_sibling("td").text.strip()
                    severity = vulnerability_soup.find("td", text="Severity").find_next_sibling("td").text.strip()
                    cvss_score = vulnerability_soup.find("td", text="CVSSv3 Score").find_next_sibling("td").text.strip()
                    impact = vulnerability_soup.find("td", text="Impact").find_next_sibling("td").text.strip()
                    cve_id = vulnerability_soup.find("td", text="CVE ID").find_next_sibling("td").text.strip()
                    summary = vulnerability_soup.find("div", class_="detail-item").find("p").text.strip()
                    timeline = vulnerability_soup.find("div", class_="detail-item", text="Timeline").find_next("p").text.strip()
                    
                    # Add the information to the Document
                    doc.add_heading(f"{ir_number} - {cve_id}", level=1)
                    doc.add_paragraph(f"Date: {date}")
                    doc.add_paragraph(f"Severity: {severity}")
                    doc.add_paragraph(f"CVSSv3 Score: {cvss_score}")
                    doc.add_paragraph(f"Impact: {impact}")
                    doc.add_paragraph(f"CVE ID: {cve_id}")
                    doc.add_paragraph(f"Summary: {summary}")
                    doc.add_paragraph(f"Timeline: {timeline}")
                    doc.add_paragraph('\n')  # Add some space between vulnerabilities
                    
                    # Increment vulnerability count
                    vulnerability_count += 1
        
        # Save the Document to a docx file with the current date in the name
        file_name = fr'C:\Users\Novo\Noticias\data\Fortiguard_{current_date}.docx'
        doc.save(file_name)
        
        # Print the number of vulnerabilities extracted
        print(f"Total number of vulnerabilities extracted: {vulnerability_count}")
    
    else:
        print("Failed to retrieve data from the URL.")

except Exception as e:
    print(f"An error occurred: {e}")
