import requests
from bs4 import BeautifulSoup
from newspaper import Article
import time
import nltk
import traceback
from docx import Document
from datetime import datetime

# Download the 'punkt' resource from NLTK
nltk.download('punkt')

url = "https://www.helpnetsecurity.com/view/news/"

# Set the User-Agent header to mimic a request from a web browser
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

try:
    # Make a request to the website with headers
    response = requests.get(url, headers=headers)
    response.raise_for_status()  # Check if the request was successful
    html = response.text

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Find the elements that contain the article titles and links
    articles = soup.find_all("div", class_="card-body")

    if not articles:
        print("No articles found on the page.")

    # Create a Document object to store the content
    doc = Document()

    # Loop through the articles and extract the information
    for article in articles:
        # Check if the <a> tag exists within the <h2> tag
        title_element = article.find("h5", class_="reset-heading card-title")
        link_element = article.find("a", class_="d-block font-size-4 pb-1")

        if title_element and link_element:
            title = title_element.text.strip()
            link = link_element["href"]

            # Create an Article object and download the content
            article_obj = Article(link)
            article_obj.download()

            # Parse the article and generate a summary
            article_obj.parse()
            article_obj.nlp()
            summary = article_obj.summary

            # Add the information to the Document
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Link: {link}")
            doc.add_paragraph(f"Summary: {summary}")
            doc.add_paragraph('\n')  # Add some space between articles

            # Introduce a delay of 2 seconds between requests
            time.sleep(2)

    # Get the current date and time
    current_date = datetime.now().strftime('%d-%m-%Y')

    # Save the Document to a docx file with the current date in the name
    file_name = fr'C:\Users\Novo\Noticias\data\Netsecurity_{current_date}.docx'
    doc.save(file_name)

except Exception as e:
    print(f"An error occurred: {e}")
    print("Exception traceback:", traceback.format_exc())
