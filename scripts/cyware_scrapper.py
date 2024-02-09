##### BUGS to Resolve
# - O modo headless aka sem pop up do web browser está bugado e da crash. (talvez por causa do chrome webdriver idk)

import requests
from bs4 import BeautifulSoup
import time
import traceback
from docx import Document
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options  # Importar as Options

# Download the 'punkt' resource from NLTK
import nltk
nltk.download('punkt')

# Set the URL
url = "https://cyware.com/cyber-security-news-articles"

# Set the User-Agent header to mimic a request from a web browser
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

# Inicializar as Options do Chrome
chrome_options = Options()

# Adicionar a opção headless (sem o pop up do browser)
#chrome_options.add_argument("--headless")

# Inicializar o Selenium WebDriver com as opções
driver = webdriver.Chrome(options=chrome_options)  # Assuming you have Chrome WebDriver installed

try:
    # Navigate to the URL
    driver.get(url)

    # Scroll down to the end of the page
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")

    # Wait for the "Load More" button to be clickable
    WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.CLASS_NAME, 'btn-primary')))

    # Click the "Load More" button
    load_more_button = driver.find_element(By.CLASS_NAME, 'btn-primary')
    load_more_button.click()

    # Wait for the additional news articles to load
    time.sleep(5)  # Adjust the sleep time based on your page load time

    # Get the HTML content of the page after clicking the button
    html = driver.page_source

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Find the elements that contain the article titles, URLs, and descriptions
    articles = soup.find_all("div", class_="cy-panel__body")

    if not articles:
        print("No articles found on the page.")

    # Create a Document object to store the content
    doc = Document()

    # Get the current date
    current_date = datetime.now().strftime('%d-%m-%Y')

    # Initialize news count
    news_count = 0

    # Loop through the articles and extract the information
    for article in articles:
        # Find the date element
        date_element = article.find("span", class_="cy-card__meta")
        if date_element:
            # Extract the date text
            article_date = date_element.text.strip()
            
            # Convert the article date string to a datetime object
            article_date = datetime.strptime(article_date, '%B %d, %Y')
            
            # Check if the article date is today's date
            if article_date.date() == datetime.now().date():
                # Find title element
                title_element = article.find("h1", class_="cy-card__title")
                if title_element:
                    title = title_element.text.strip()

                    # Find all <a> tags containing target="_blank"
                    link_element = article.find("a", target="_blank")
                    if link_element:
                        link = link_element["href"]
                        
                        # Check if the domain is missing in the URL
                        if not link.startswith("http"):
                            # Add domain if missing
                            link = "https://cyware.com" + link
                            
                    # Find description element
                    description_element = article.find("div", class_="cy-card__description")
                    if description_element:
                        description = description_element.text.strip()

                        # Add the information to the Document
                        doc.add_heading(title, level=1)
                        doc.add_paragraph(f"Link: {link}")
                        doc.add_paragraph(f"Description: {description}")
                        doc.add_paragraph('\n')  # Add some space between articles
                        
                        # Increment news count
                        news_count += 1

    # Save the Document to a docx file with the current date in the name
    file_name = fr'C:\Users\Novo\Noticias\data\Cyware_{current_date}.docx'
    doc.save(file_name)
    
    # Print the number of news inserted
    print(f"[Cyware] Número Total de Notícias novas: {news_count}")

except Exception as e:
    print(f"An error occurred: {e}")
    print("Exception traceback:", traceback.format_exc())

finally:
    # Quit the WebDriver
    driver.quit()
